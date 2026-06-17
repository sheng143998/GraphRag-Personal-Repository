import asyncio
import base64
import io
import os

import pytest

os.environ["AI_RAG_USE_DATABASE"] = "false"
os.environ["MODEL_PROVIDER"] = "stub"
os.environ["LLM_PROVIDER"] = "stub"
os.environ["EMBEDDING_PROVIDER"] = "stub"
os.environ["RERANK_PROVIDER"] = "stub"

from app.core.constants import DocumentType, FileType
from app.rag.parsers.base import DocxParser
from app.schemas.ingest import DocumentIngestRequest, DocumentPayload


def test_docx_parser_preserves_heading_and_table_structure() -> None:
    parsed = asyncio.run(_parse_structured_docx())

    assert "# Architecture" in parsed.text
    assert "## Retrieval" in parsed.text
    assert "| metric | meaning |" in parsed.text
    assert "| --- | --- |" in parsed.text
    assert "| recall | hit expected evidence |" in parsed.text
    assert "Table 1 Rows" in parsed.text
    assert "Row 2: R2C1=recall | R2C2=hit expected evidence" in parsed.text
    assert parsed.metadata["heading_count"] == 2
    assert parsed.metadata["table_count"] == 1


def test_docx_parser_preserves_nested_and_merged_table_context() -> None:
    parsed = asyncio.run(_parse_complex_table_docx())

    assert "Nested table" in parsed.text
    assert "R2C2=Nested table 1: Row 1: R1C1=e | R1C2=f ; Row 2: R2C1=g&t | R2C2=h" in parsed.text
    assert "R1C2=b (continues into C3)" in parsed.text
    assert "R2C2=c | R2C3=d" in parsed.text
    assert parsed.metadata["table_count"] == 2


def test_docx_parser_reports_invalid_docx_content() -> None:
    with pytest.raises(RuntimeError, match="docx-parse-failed"):
        asyncio.run(_parse_invalid_docx())


async def _parse_structured_docx():
    return await DocxParser().parse(
        raw_content="garbage text that should not be used when content_base64 exists",
        request=DocumentIngestRequest(
            knowledge_base_id="kb-docx",
            document_id="doc-docx",
            title="DOCX Notes",
            document_type=DocumentType.TECH_NOTE,
            file=DocumentPayload(
                filename="notes.docx",
                file_type=FileType.DOCX,
                content_base64=_structured_docx_base64(),
            ),
        ),
    )


async def _parse_complex_table_docx():
    return await DocxParser().parse(
        raw_content="",
        request=DocumentIngestRequest(
            knowledge_base_id="kb-docx",
            document_id="doc-complex-docx",
            title="Complex DOCX Tables",
            document_type=DocumentType.TECH_NOTE,
            file=DocumentPayload(
                filename="complex-tables.docx",
                file_type=FileType.DOCX,
                content_base64=_complex_table_docx_base64(),
            ),
        ),
    )


async def _parse_invalid_docx():
    return await DocxParser().parse(
        raw_content="",
        request=DocumentIngestRequest(
            knowledge_base_id="kb-docx",
            document_id="doc-invalid-docx",
            title="Broken DOCX",
            document_type=DocumentType.TECH_NOTE,
            file=DocumentPayload(
                filename="broken.docx",
                file_type=FileType.DOCX,
                content_base64=base64.b64encode(b"not a real docx zip").decode("ascii"),
            ),
        ),
    )


def _structured_docx_base64() -> str:
    import docx

    document = docx.Document()
    document.add_heading("Architecture", level=1)
    document.add_paragraph("The document keeps Word heading structure.")
    document.add_heading("Retrieval", level=2)
    table = document.add_table(rows=3, cols=2)
    table.cell(0, 0).text = "metric"
    table.cell(0, 1).text = "meaning"
    table.cell(1, 0).text = "recall"
    table.cell(1, 1).text = "hit expected evidence"
    table.cell(2, 0).text = "precision"
    table.cell(2, 1).text = "avoid irrelevant chunks"

    buffer = io.BytesIO()
    document.save(buffer)
    return base64.b64encode(buffer.getvalue()).decode("ascii")


def _complex_table_docx_base64() -> str:
    import docx

    document = docx.Document()
    document.add_paragraph("Nested Table")
    table = document.add_table(rows=3, cols=3)
    table.cell(0, 0).text = "a"
    table.cell(0, 1).text = ">b<"
    table.cell(0, 2).text = "c"
    table.cell(1, 0).text = "d"
    table.cell(1, 2).text = "i"
    nested = table.cell(1, 1).add_table(rows=2, cols=2)
    nested.cell(0, 0).text = "e"
    nested.cell(0, 1).text = "f"
    nested.cell(1, 0).text = "g&t"
    nested.cell(1, 1).text = "h"
    table.cell(2, 0).text = "j"
    table.cell(2, 1).text = "k"
    table.cell(2, 2).text = "l"

    document.add_paragraph("Table with merged cells")
    merged = document.add_table(rows=3, cols=3)
    merged.cell(0, 0).text = "a"
    merged.cell(0, 1).text = "b"
    merged.cell(0, 1).merge(merged.cell(0, 2))
    merged.cell(1, 1).text = "c"
    merged.cell(1, 2).text = "d"
    merged.cell(2, 0).text = "e"
    merged.cell(2, 0).merge(merged.cell(2, 1))

    buffer = io.BytesIO()
    document.save(buffer)
    return base64.b64encode(buffer.getvalue()).decode("ascii")
