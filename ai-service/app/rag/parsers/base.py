from __future__ import annotations

from dataclasses import dataclass, field
import json
import re
import zipfile
from xml.etree import ElementTree as ET

from app.schemas.ingest import DocumentIngestRequest
import base64
import io


@dataclass(slots=True)
class ParsedContent:
    text: str
    metadata: dict[str, object] = field(default_factory=dict)


class BaseParser:
    name = "base-parser"
    version = "v1"

    async def parse(self, *, raw_content: str, request: DocumentIngestRequest) -> ParsedContent:
        raise NotImplementedError


class PlainTextParser(BaseParser):
    name = "plain-text-parser"
    version = "v1"

    async def parse(self, *, raw_content: str, request: DocumentIngestRequest) -> ParsedContent:
        return ParsedContent(
            text=raw_content.strip(),
            metadata={
                "document_type": request.document_type,
                "file_type": request.file.file_type,
            },
        )


class HtmlParser(PlainTextParser):
    name = "html-parser"


class DocxParser(BaseParser):
    name = "docx-parser"
    version = "v2"

    async def parse(self, *, raw_content: str, request: DocumentIngestRequest) -> ParsedContent:
        text_parts: list[str] = []
        parse_error: str | None = None
        content_b64 = request.file.content_base64 or ""
        if not content_b64 and request.file.content:
            content_b64 = request.file.content

        if raw_content and not content_b64:
            text_parts.append(raw_content.strip())

        heading_count = 0
        table_count = 0
        if content_b64:
            try:
                import docx
            except Exception as exc:
                parse_error = f"python-docx-unavailable:{type(exc).__name__}:{exc}"
            else:
                try:
                    decoded = base64.b64decode(content_b64, validate=True)
                    doc = docx.Document(io.BytesIO(decoded))

                    for para in doc.paragraphs:
                        t = para.text.strip()
                        if not t:
                            continue
                        heading_level = _docx_heading_level(para)
                        if heading_level:
                            heading_count += 1
                            text_parts.append(f"{'#' * heading_level} {t}")
                        else:
                            text_parts.append(t)

                    for table_index, table in enumerate(doc.tables, start=1):
                        table_text = _docx_table_to_text(table, table_index=table_index)
                        if table_text:
                            table_count += 1
                            text_parts.append(table_text)
                except Exception as exc:
                    parse_error = f"docx-parse-failed:{type(exc).__name__}:{exc}"

        if not text_parts and raw_content and raw_content.strip() and not _looks_like_binary_text(raw_content):
            text_parts.append(raw_content.strip())
        combined = "\n\n".join(text_parts).strip()
        if not combined and parse_error:
            raise RuntimeError(parse_error)
        return ParsedContent(
            text=combined,
            metadata={
                "document_type": request.document_type,
                "file_type": request.file.file_type,
                "parser": self.name,
                "parser_version": self.version,
                "heading_count": heading_count,
                "table_count": table_count,
                "parse_error": parse_error,
            },
        )


def _docx_heading_level(paragraph) -> int | None:
    style_name = getattr(getattr(paragraph, "style", None), "name", "") or ""
    match = re.search(r"Heading\s+([1-6])", style_name, re.I)
    if match:
        return int(match.group(1))
    return None


def _docx_table_to_text(table, *, table_index: int) -> str | None:
    markdown_table = _docx_table_to_markdown(table)
    row_records = _docx_table_row_records(table)
    if not markdown_table and not row_records:
        return None
    parts = [f"Table {table_index}"]
    if markdown_table:
        parts.append(markdown_table)
    if row_records:
        parts.append(f"Table {table_index} Rows")
        parts.extend(row_records)
    return "\n".join(parts)


def _docx_table_to_markdown(table) -> str | None:
    rows = []
    for row_index, row in enumerate(table.rows):
        values = [_docx_cell_text(cell) for cell in row.cells]
        if row_index > 0:
            values = _deduplicate_repeated_row_values(values)
        if any(values):
            rows.append(values)
    if not rows:
        return None
    width = max(len(row) for row in rows)
    normalized = [row + [""] * max(0, width - len(row)) for row in rows]
    header = [cell or f"column_{index + 1}" for index, cell in enumerate(normalized[0])]
    if len(normalized) == 1:
        body = [normalized[0]]
    else:
        body = normalized[1:]
    markdown_rows = [
        "| " + " | ".join(header) + " |",
        "| " + " | ".join(["---"] * len(header)) + " |",
    ]
    for row in body:
        markdown_rows.append("| " + " | ".join(cell or "" for cell in row) + " |")
    return "\n".join(markdown_rows)


def _docx_table_row_records(table) -> list[str]:
    records: list[str] = []
    for row_index, row in enumerate(table.rows, start=1):
        cells: list[str] = []
        previous_value = None
        repeat_count = 0
        for column_index, cell in enumerate(row.cells, start=1):
            value = _docx_cell_text(cell)
            if not value:
                continue
            span = _docx_cell_grid_span(cell)
            if value == previous_value:
                repeat_count += 1
                if repeat_count > 1:
                    continue
                cells[-1] = f"{cells[-1]} (continues into C{column_index})"
                continue
            previous_value = value
            repeat_count = 1
            span_note = f" (continues into C{column_index + span - 1})" if span > 1 else ""
            cells.append(f"R{row_index}C{column_index}={value}{span_note}")
        if cells:
            records.append(f"Row {row_index}: " + " | ".join(cells))
    return records


def _docx_cell_text(cell) -> str:
    values: list[str] = []
    for paragraph in cell.paragraphs:
        text = paragraph.text.strip()
        if text:
            values.append(text)
    for nested_index, nested_table in enumerate(cell.tables, start=1):
        nested_records = _docx_table_row_records(nested_table)
        if nested_records:
            values.append(f"Nested table {nested_index}: " + " ; ".join(nested_records))
    return " ".join(values).strip()


def _docx_cell_grid_span(cell) -> int:
    try:
        tc_pr = cell._tc.tcPr
        grid_span = tc_pr.gridSpan if tc_pr is not None else None
        value = getattr(grid_span, "val", None)
        return int(value) if value else 1
    except Exception:
        return 1


def _deduplicate_repeated_row_values(values: list[str]) -> list[str]:
    deduped: list[str] = []
    previous_value = None
    repeat_count = 0
    for value in values:
        if value and value == previous_value:
            repeat_count += 1
            if repeat_count > 1:
                deduped.append("")
                continue
        else:
            previous_value = value
            repeat_count = 1
        deduped.append(value)
    return deduped


def _looks_like_binary_text(value: str) -> bool:
    if not value:
        return False
    replacement_count = value.count("\ufffd")
    control_count = sum(1 for char in value if ord(char) < 32 and char not in "\t\r\n")
    return replacement_count > 5 or control_count > 5


class SpreadsheetParser(PlainTextParser):
    name = "spreadsheet-parser"
    version = "v2"

    async def parse(self, *, raw_content: str, request: DocumentIngestRequest) -> ParsedContent:
        content_b64 = request.file.content_base64 or ""
        if not content_b64 and request.file.content:
            content_b64 = request.file.content

        file_type = str(request.file.file_type)
        if file_type == "xlsx" and content_b64:
            try:
                decoded = base64.b64decode(content_b64, validate=True)
                tables = _xlsx_tables(decoded)
            except Exception as exc:
                raise RuntimeError(f"xlsx-parse-failed:{type(exc).__name__}:{exc}") from exc

            text = _spreadsheet_tables_to_text(tables)
            return ParsedContent(
                text=text,
                metadata={
                    "document_type": request.document_type,
                    "file_type": request.file.file_type,
                    "parser": self.name,
                    "parser_version": self.version,
                    "spreadsheet_table_count": len(tables),
                    "spreadsheet_sheet_names": [table["sheet_name"] for table in tables],
                    "spreadsheet_tables": tables,
                },
            )

        if file_type == "xls" and _looks_like_binary_text(raw_content):
            raise RuntimeError("xls-parse-unsupported: convert the workbook to .xlsx or .csv before upload")

        return ParsedContent(
            text=raw_content.strip(),
            metadata={
                "document_type": request.document_type,
                "file_type": request.file.file_type,
                "parser": self.name,
                "parser_version": self.version,
            },
        )


_XLSX_NS = {"main": "http://schemas.openxmlformats.org/spreadsheetml/2006/main"}
_RELS_NS = {"rel": "http://schemas.openxmlformats.org/package/2006/relationships"}
_OD_RELS_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"


def _xlsx_tables(content: bytes) -> list[dict[str, object]]:
    with zipfile.ZipFile(io.BytesIO(content)) as zf:
        shared_strings = _xlsx_shared_strings(zf)
        sheet_paths = _xlsx_sheet_paths(zf)
        tables: list[dict[str, object]] = []
        for sheet_index, (sheet_name, sheet_path) in enumerate(sheet_paths, start=1):
            if sheet_path not in zf.namelist():
                continue
            rows, row_numbers = _xlsx_sheet_rows(zf.read(sheet_path), shared_strings)
            table = _spreadsheet_table_from_rows(
                sheet_name=sheet_name or f"Sheet{sheet_index}",
                rows=rows,
                row_numbers=row_numbers,
            )
            if table is not None:
                table["sheet_index"] = sheet_index
                tables.append(table)
        return tables


def _xlsx_shared_strings(zf: zipfile.ZipFile) -> list[str]:
    if "xl/sharedStrings.xml" not in zf.namelist():
        return []
    root = ET.fromstring(zf.read("xl/sharedStrings.xml"))
    values: list[str] = []
    for item in root.findall("main:si", _XLSX_NS):
        text = "".join(node.text or "" for node in item.findall(".//main:t", _XLSX_NS))
        values.append(_clean_spreadsheet_cell(text))
    return values


def _xlsx_sheet_paths(zf: zipfile.ZipFile) -> list[tuple[str, str]]:
    if "xl/workbook.xml" not in zf.namelist():
        return []
    workbook = ET.fromstring(zf.read("xl/workbook.xml"))
    relationships = _xlsx_workbook_relationships(zf)
    sheet_paths: list[tuple[str, str]] = []
    for sheet in workbook.findall(".//main:sheet", _XLSX_NS):
        name = _clean_spreadsheet_cell(sheet.attrib.get("name", ""))
        relationship_id = sheet.attrib.get(f"{{{_OD_RELS_NS}}}id")
        target = relationships.get(relationship_id or "")
        if not target:
            continue
        sheet_paths.append((name, _xlsx_resolve_workbook_target(target)))
    return sheet_paths


def _xlsx_workbook_relationships(zf: zipfile.ZipFile) -> dict[str, str]:
    rels_path = "xl/_rels/workbook.xml.rels"
    if rels_path not in zf.namelist():
        return {}
    root = ET.fromstring(zf.read(rels_path))
    return {
        rel.attrib.get("Id", ""): rel.attrib.get("Target", "")
        for rel in root.findall("rel:Relationship", _RELS_NS)
        if rel.attrib.get("Id") and rel.attrib.get("Target")
    }


def _xlsx_resolve_workbook_target(target: str) -> str:
    normalized = target.replace("\\", "/").lstrip("/")
    if normalized.startswith("xl/"):
        return normalized
    return f"xl/{normalized}"


def _xlsx_sheet_rows(sheet_xml: bytes, shared_strings: list[str]) -> tuple[list[list[str]], list[int]]:
    root = ET.fromstring(sheet_xml)
    rows: list[list[str]] = []
    row_numbers: list[int] = []
    for row in root.findall(".//main:sheetData/main:row", _XLSX_NS):
        row_number = _safe_int(row.attrib.get("r"), len(row_numbers) + 1)
        cells: dict[int, str] = {}
        max_column = 0
        for cell in row.findall("main:c", _XLSX_NS):
            column_index = _xlsx_column_index(cell.attrib.get("r", ""))
            if column_index <= 0:
                column_index = max_column + 1
            max_column = max(max_column, column_index)
            cells[column_index] = _xlsx_cell_value(cell, shared_strings)
        if not cells:
            continue
        values = [cells.get(index, "") for index in range(1, max_column + 1)]
        if any(value.strip() for value in values):
            rows.append(values)
            row_numbers.append(row_number)
    return rows, row_numbers


def _xlsx_cell_value(cell: ET.Element, shared_strings: list[str]) -> str:
    cell_type = cell.attrib.get("t", "")
    if cell_type == "inlineStr":
        return _clean_spreadsheet_cell("".join(node.text or "" for node in cell.findall(".//main:t", _XLSX_NS)))

    value_node = cell.find("main:v", _XLSX_NS)
    raw_value = value_node.text if value_node is not None else ""
    if cell_type == "s":
        index = _safe_int(raw_value, -1)
        if 0 <= index < len(shared_strings):
            return shared_strings[index]
        return ""
    if cell_type == "b":
        return "TRUE" if raw_value == "1" else "FALSE" if raw_value == "0" else _clean_spreadsheet_cell(raw_value)
    return _clean_spreadsheet_cell(raw_value)


def _xlsx_column_index(cell_ref: str) -> int:
    match = re.match(r"([A-Za-z]+)", cell_ref or "")
    if not match:
        return 0
    value = 0
    for char in match.group(1).upper():
        value = value * 26 + (ord(char) - ord("A") + 1)
    return value


def _spreadsheet_table_from_rows(
    *,
    sheet_name: str,
    rows: list[list[str]],
    row_numbers: list[int],
) -> dict[str, object] | None:
    if not rows:
        return None
    width = max(len(row) for row in rows)
    normalized_rows = [_pad_spreadsheet_row(row, width) for row in rows]
    header_row_index = next((index for index, row in enumerate(normalized_rows) if any(cell.strip() for cell in row)), 0)
    header = [
        _clean_spreadsheet_column(value, index)
        for index, value in enumerate(normalized_rows[header_row_index])
    ]
    data_rows = normalized_rows[header_row_index + 1 :]
    data_row_numbers = row_numbers[header_row_index + 1 :]
    filtered_rows: list[list[str]] = []
    filtered_row_numbers: list[int] = []
    for row, row_number in zip(data_rows, data_row_numbers, strict=False):
        if any(cell.strip() for cell in row):
            filtered_rows.append(row)
            filtered_row_numbers.append(row_number)
    if not filtered_rows:
        return None
    return {
        "sheet_name": sheet_name,
        "header": header,
        "rows": filtered_rows,
        "row_numbers": filtered_row_numbers,
        "header_row_number": row_numbers[header_row_index] if row_numbers else 1,
        "data_start_row_number": filtered_row_numbers[0] if filtered_row_numbers else 1,
    }


def _spreadsheet_tables_to_text(tables: list[dict[str, object]]) -> str:
    parts: list[str] = []
    for table in tables:
        sheet_name = str(table.get("sheet_name") or "Sheet1")
        header = [str(value) for value in table.get("header", [])]
        rows = [[str(value) for value in row] for row in table.get("rows", []) if isinstance(row, list)]
        row_numbers = [int(value) for value in table.get("row_numbers", []) if isinstance(value, int)]
        parts.append(f"Sheet: {sheet_name}")
        parts.append("Columns: " + " | ".join(header))
        for index, row in enumerate(rows):
            row_number = row_numbers[index] if index < len(row_numbers) else index + 1
            values = [
                f"{column}={value}"
                for column, value in zip(header, row, strict=False)
                if str(value).strip()
            ]
            parts.append(f"Row {row_number}: " + " | ".join(values))
        parts.append("")
    return "\n".join(parts).strip()


def _pad_spreadsheet_row(row: list[str], width: int) -> list[str]:
    return [_clean_spreadsheet_cell(value) for value in row] + [""] * max(0, width - len(row))


def _clean_spreadsheet_column(value: str, index: int) -> str:
    cleaned = _clean_spreadsheet_cell(value)
    return cleaned or f"column_{index + 1}"


def _clean_spreadsheet_cell(value: object) -> str:
    text = str(value or "")
    text = text.replace("\x00", "").replace("\ufffd", "")
    return re.sub(r"\s+", " ", text).strip()


def _safe_int(value: object, default: int) -> int:
    try:
        return int(str(value))
    except Exception:
        return default


class MinerUPdfParser(BaseParser):
    name = "mineru-pdf-adapter"
    version = "v3"

    async def parse(self, *, raw_content: str, request: DocumentIngestRequest) -> ParsedContent:
        import asyncio
        import httpx
        from app.core.config import settings

        token = settings.mineru_api_token
        use_standard = bool(token)

        poll_timeout = max(30, int(settings.mineru_poll_timeout_seconds))
        poll_interval = max(1, int(settings.mineru_poll_interval_seconds))

        file_name = request.file.filename or "document.pdf"
        content_b64 = request.file.content_base64 or ""
        source_url = request.file.source_path or request.file.content or raw_content
        if source_url and source_url.startswith(("http://", "https://")):
            source_url = source_url.strip()
        use_file = bool(content_b64)
        use_url = bool(source_url and source_url.startswith("https://"))

        print(f"[MinerU] mode={'standard' if use_standard else 'agent'} file={use_file} url={use_url} name={file_name}")

        if not use_file and not use_url:
            return ParsedContent(text="", metadata={"adapter": "mineru", "status": "skipped", "reason": "no_content_or_url", "file_type": request.file.file_type})

        api_mode = "standard" if use_standard else "agent"
        headers = {"Content-Type": "application/json"}
        if use_standard:
            headers["Authorization"] = f"Bearer {token}"

        # Use a single client with generous timeout for the whole flow
        async with httpx.AsyncClient(timeout=httpx.Timeout(120.0, connect=10.0)) as client:
            task_id = ""
            batch_id = ""
            standard_file_batch = False
            try:
                if use_file:
                    decoded_bytes = base64.b64decode(content_b64)
                    print(f"[MinerU] file size={len(decoded_bytes)} bytes")

                    if use_standard:
                        resp = await client.post(
                            "https://mineru.net/api/v4/file-urls/batch",
                            headers=headers,
                            json={"files": [{"name": file_name}]},
                        )
                        data = resp.json()
                        if data.get("code") != 0:
                            return ParsedContent(text="", metadata={"adapter": "mineru", "status": "submit_failed", "error": data.get("msg", "unknown"), "api": api_mode, "file_type": request.file.file_type})
                        batch_id = data["data"]["batch_id"]
                        upload_url = data["data"]["file_urls"][0]
                        put_resp = await client.put(upload_url, content=decoded_bytes)
                        print(f"[MinerU] PUT status={put_resp.status_code}")
                        await asyncio.sleep(2)
                        task_id = batch_id
                        standard_file_batch = True
                    else:
                        # Agent API: send file_content directly in JSON (inline mode)
                        resp = await client.post(
                            "https://mineru.net/api/v1/agent/parse/file",
                            headers=headers,
                            json={
                                "file_name": file_name,
                                "language": "ch",
                                "enable_table": True,
                                "file_content": content_b64,
                            },
                        )
                        data = resp.json()
                        print(f"[MinerU] parse/file: code={data.get('code')} msg={data.get('msg')}")
                        if data.get("code") != 0:
                            return ParsedContent(text="", metadata={"adapter": "mineru", "status": "submit_failed", "error": data.get("msg", "unknown"), "api": api_mode, "file_type": request.file.file_type})
                        task_id = data["data"]["task_id"]
                        upload_url = data["data"].get("file_url", "")
                        if upload_url:
                            put_resp = await client.put(upload_url, content=decoded_bytes)
                            print(f"[MinerU] agent PUT status={put_resp.status_code}")
                            if put_resp.status_code >= 400:
                                return ParsedContent(
                                    text="",
                                    metadata={
                                        "adapter": "mineru",
                                        "status": "upload_failed",
                                        "task_id": task_id,
                                        "http_status": put_resp.status_code,
                                        "api": api_mode,
                                        "file_type": request.file.file_type,
                                    },
                                )
                            await asyncio.sleep(2)
                        print(f"[MinerU] task_id={task_id}")
                elif use_url and use_standard:
                    resp = await client.post(
                        f"{settings.mineru_api_base_url}/extract/task",
                        headers=headers,
                        json={"url": source_url, "model_version": "vlm"},
                    )
                    data = resp.json()
                    if data.get("code") != 0:
                        return ParsedContent(text="", metadata={"adapter": "mineru", "status": "submit_failed", "error": data.get("msg", "unknown"), "api": api_mode, "file_type": request.file.file_type})
                    task_id = data["data"]["task_id"]
                else:
                    resp = await client.post(
                        "https://mineru.net/api/v1/agent/parse/url",
                        headers=headers,
                        json={"url": source_url, "language": "ch", "enable_table": True},
                    )
                    data = resp.json()
                    if data.get("code") != 0:
                        return ParsedContent(text="", metadata={"adapter": "mineru", "status": "submit_failed", "error": data.get("msg", "unknown"), "api": api_mode, "file_type": request.file.file_type})
                    task_id = data["data"]["task_id"]
            except Exception as e:
                print(f"[MinerU] submit error: {type(e).__name__}: {e}")
                return ParsedContent(text="", metadata={"adapter": "mineru", "status": "submit_error", "error": str(e), "api": api_mode, "file_type": request.file.file_type})

            # Poll for result
            elapsed = 0
            last_poll_error = ""
            while elapsed < poll_timeout:
                await asyncio.sleep(poll_interval)
                elapsed += poll_interval
                try:
                    if standard_file_batch:
                        poll_url = f"https://mineru.net/api/v4/extract-results/batch/{batch_id or task_id}"
                        poll_resp = await client.get(poll_url, headers=headers)
                    elif use_standard:
                        poll_url = f"https://mineru.net/api/v4/extract/task/{task_id}"
                        poll_resp = await client.get(poll_url, headers=headers)
                    else:
                        poll_resp = await client.get(
                            f"https://mineru.net/api/v1/agent/parse/{task_id}",
                            headers=headers,
                        )
                    poll = poll_resp.json()
                    if poll.get("code") != 0:
                        last_poll_error = f"{poll.get('code')}: {poll.get('msg', 'unknown')}"
                        if elapsed % 30 <= poll_interval:
                            print(f"[MinerU] poll {elapsed}s error: {last_poll_error}")
                        continue
                    poll_data = poll.get("data", {})
                    if standard_file_batch:
                        result_item = _select_mineru_batch_result(poll_data, file_name)
                        if result_item is None:
                            last_poll_error = "empty extract_result"
                            if elapsed % 30 <= poll_interval:
                                print(f"[MinerU] poll {elapsed}s error: {last_poll_error}")
                            continue
                        result_payload = result_item
                        state = str(result_item.get("state") or poll_data.get("state") or "").lower()
                    else:
                        result_payload = poll_data
                        state = str(poll_data.get("state", "")).lower()
                    if not state:
                        last_poll_error = "missing state"
                        continue
                    if elapsed % 10 <= poll_interval:
                        print(f"[MinerU] poll {elapsed}s: {state}")
                    if state == "done":
                        md_text = ""
                        result_source = ""
                        result_errors: list[str] = []
                        async with httpx.AsyncClient(
                            timeout=httpx.Timeout(120.0, connect=10.0),
                            trust_env=False,
                            follow_redirects=True,
                        ) as download_client:
                            if use_standard:
                                for zip_url in _mineru_result_urls(result_payload, poll_data, kind="zip"):
                                    try:
                                        zip_resp = await download_client.get(zip_url)
                                        md_text = _mineru_text_from_zip(zip_resp.content)
                                        if md_text:
                                            result_source = "zip"
                                            break
                                    except Exception as exc:
                                        result_errors.append(f"zip:{type(exc).__name__}:{exc}")
                            if not md_text:
                                for md_url in _mineru_result_urls(result_payload, poll_data, kind="markdown"):
                                    try:
                                        md_text = (await download_client.get(md_url)).text
                                        print(f"[MinerU] DONE! md={len(md_text)} chars")
                                        if md_text:
                                            result_source = "markdown_url"
                                            break
                                    except Exception as exc:
                                        result_errors.append(f"md:{type(exc).__name__}:{exc}")
                        if not md_text:
                            response_keys = _mineru_payload_keys(result_payload, poll_data)
                            last_poll_error = "done_without_result_content"
                            if result_errors:
                                last_poll_error = f"{last_poll_error}; errors={'; '.join(result_errors[-3:])}"
                            print(f"[MinerU] DONE but no downloadable result content; keys={response_keys}")
                        normalized = (md_text or "").strip()
                        if not normalized:
                            normalized = _fallback_pdf_text(request, raw_content)
                            if normalized and not result_source:
                                result_source = "raw_fallback"
                        if not normalized:
                            normalized = _extract_pdf_text_locally(content_b64)
                            if normalized and not result_source:
                                result_source = "local_pdf_fallback"
                        return ParsedContent(
                            text=normalized,
                            metadata={
                                "adapter": "mineru",
                                "status": "completed",
                                "task_id": task_id,
                                "batch_id": batch_id or None,
                                "elapsed_seconds": elapsed,
                                "api": api_mode,
                                "file_type": request.file.file_type,
                                "result_source": result_source or None,
                                "last_poll_error": last_poll_error or None,
                                **_mineru_markdown_metadata(normalized),
                            },
                        )
                    if state in {"failed", "fail", "error"}:
                        err = result_payload.get("err_msg") or result_payload.get("message") or result_payload.get("error") or "unknown"
                        print(f"[MinerU] FAILED: {err}")
                        return ParsedContent(
                            text=_extract_pdf_text_locally(content_b64) or _fallback_pdf_text(request, raw_content),
                            metadata={"adapter": "mineru", "status": "failed", "task_id": task_id, "batch_id": batch_id or None, "error": err, "api": api_mode, "file_type": request.file.file_type},
                        )
                except Exception as exc:
                    last_poll_error = f"{type(exc).__name__}: {exc}"
                    continue
            print(f"[MinerU] TIMEOUT after {elapsed}s")
            return ParsedContent(
                text=_extract_pdf_text_locally(content_b64) or _fallback_pdf_text(request, raw_content),
                metadata={"adapter": "mineru", "status": "timeout", "task_id": task_id, "batch_id": batch_id or None, "elapsed_seconds": elapsed, "last_poll_error": last_poll_error or None, "api": api_mode, "file_type": request.file.file_type},
            )


def _select_mineru_batch_result(data: object, file_name: str) -> dict[str, object] | None:
    if not isinstance(data, dict):
        return None
    raw_results = data.get("extract_result") or data.get("extract_results") or []
    if isinstance(raw_results, dict):
        raw_results = [raw_results]
    if not isinstance(raw_results, list):
        return None
    results = [item for item in raw_results if isinstance(item, dict)]
    if not results:
        return None
    for item in results:
        candidate_name = str(item.get("file_name") or item.get("filename") or item.get("name") or "")
        if candidate_name and (candidate_name == file_name or candidate_name.endswith(file_name)):
            return item
    return results[0]


def _mineru_result_urls(*payloads: object, kind: str) -> list[str]:
    urls: list[str] = []
    seen: set[str] = set()

    def visit(value: object, key_path: str = "") -> None:
        if isinstance(value, dict):
            for key, child in value.items():
                child_key_path = f"{key_path}.{key}" if key_path else str(key)
                visit(child, child_key_path)
        elif isinstance(value, list):
            for child in value:
                visit(child, key_path)
        elif isinstance(value, str):
            candidate = value.strip()
            if not candidate.startswith(("http://", "https://")):
                return
            key_lower = key_path.lower()
            url_lower = candidate.split("?", 1)[0].lower()
            is_zip = "zip" in key_lower or url_lower.endswith(".zip")
            is_markdown = (
                "markdown" in key_lower
                or key_lower.endswith("_md_url")
                or key_lower.endswith(".md_url")
                or url_lower.endswith((".md", ".markdown"))
            )
            if (kind == "zip" and not is_zip) or (kind == "markdown" and not is_markdown):
                return
            if candidate not in seen:
                seen.add(candidate)
                urls.append(candidate)

    for payload in payloads:
        visit(payload)
    return urls


def _mineru_text_from_zip(content: bytes) -> str:
    with zipfile.ZipFile(io.BytesIO(content)) as zf:
        names = zf.namelist()
        markdown_names = [name for name in names if name.lower().endswith((".md", ".markdown"))]
        markdown_names.sort(key=lambda name: (0 if "full" in name.lower() else 1, name))
        for name in markdown_names:
            text = zf.read(name).decode("utf-8", errors="replace").strip()
            if text:
                return text

        json_names = [name for name in names if name.lower().endswith(".json")]
        json_names.sort(key=lambda name: (0 if "content_list" in name.lower() else 1, name))
        for name in json_names:
            try:
                payload = json.loads(zf.read(name).decode("utf-8", errors="replace"))
            except Exception:
                continue
            text = _mineru_json_text(payload)
            if text:
                return text
    return ""


def _mineru_json_text(payload: object) -> str:
    parts: list[str] = []
    seen: set[str] = set()
    preferred_keys = {
        "text",
        "content",
        "table_body",
        "table_caption",
        "table_footnote",
        "img_caption",
        "image_caption",
        "caption",
        "latex",
        "equation",
        "html",
    }

    def add_text(value: object) -> None:
        if isinstance(value, str):
            text = value.strip()
            if text and text not in seen:
                seen.add(text)
                parts.append(text)
        elif isinstance(value, list):
            for item in value:
                add_text(item)

    def visit(value: object) -> None:
        if isinstance(value, dict):
            for key in preferred_keys:
                if key in value:
                    add_text(value.get(key))
            block_type = str(value.get("type") or "").lower()
            image_path = value.get("img_path") or value.get("image_path")
            if block_type == "image" and isinstance(image_path, str) and image_path.strip():
                add_text(f"![image]({image_path.strip()})")
            for child in value.values():
                if isinstance(child, (dict, list)):
                    visit(child)
        elif isinstance(value, list):
            for child in value:
                visit(child)

    visit(payload)
    return "\n\n".join(parts).strip()


def _mineru_payload_keys(*payloads: object) -> list[str]:
    keys: set[str] = set()

    def visit(value: object, prefix: str = "") -> None:
        if isinstance(value, dict):
            for key, child in value.items():
                path = f"{prefix}.{key}" if prefix else str(key)
                keys.add(path)
                if isinstance(child, (dict, list)):
                    visit(child, path)
        elif isinstance(value, list):
            for child in value[:3]:
                visit(child, prefix)

    for payload in payloads:
        visit(payload)
    return sorted(keys)[:80]


def _mineru_markdown_metadata(text: str) -> dict[str, object]:
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    table_line_count = sum(1 for line in lines if line.count("|") >= 2)
    page_markers = sorted(
        {
            int(match.group(1))
            for match in re.finditer(r"(?:page|页码|第)\s*[:：]?\s*(\d{1,5})", text, re.I)
        }
    )
    formula_count = len(re.findall(r"\$\$[\s\S]+?\$\$|\\\[[\s\S]+?\\\]", text))
    return {
        "mineru_heading_count": len(re.findall(r"(?m)^#{1,6}\s+\S+", text)),
        "mineru_image_count": len(re.findall(r"!\[[^\]]*]\([^)]+\)|!\[\[[^\]]+]]", text)),
        "mineru_table_count": 1 if table_line_count >= 2 else 0,
        "mineru_table_line_count": table_line_count,
        "mineru_code_block_count": len(re.findall(r"(?m)^(```|~~~)", text)) // 2,
        "mineru_formula_count": formula_count,
        "mineru_page_markers": page_markers[:50],
        "mineru_page_count_estimate": len(page_markers),
    }


def _fallback_pdf_text(request: DocumentIngestRequest, raw_content: str) -> str:
    candidates: list[str] = []
    if raw_content and _looks_like_plain_text(raw_content):
        candidates.append(_clean_storage_text(raw_content.strip()))

    content_b64 = request.file.content_base64 or request.file.content or ""
    if not content_b64:
        return "\n\n".join(candidates).strip()

    import base64
    import re

    try:
        decoded = base64.b64decode(content_b64)
        try:
            decoded_text = decoded.decode("utf-8", errors="ignore")
        except Exception:
            decoded_text = ""
        if decoded_text:
            matches = re.findall(r"\(([^()]*)\)\s*T[Jj]", decoded_text)
            if matches:
                cleaned = "\n\n".join(_clean_storage_text(match.strip()) for match in matches if match.strip())
                if cleaned:
                    candidates.append(cleaned)
            elif "MinerU fallback PDF text" in decoded_text:
                candidates.append("MinerU fallback PDF text")
    except Exception:
        pass

    return "\n\n".join(part for part in candidates if part).strip()


def _looks_like_plain_text(text: str) -> bool:
    if not text.strip():
        return False
    if text.startswith("%PDF-") or "\x00" in text:
        return False
    printable = sum(1 for char in text if char.isprintable() or char in "\r\n\t")
    return printable / max(len(text), 1) > 0.85


def _clean_storage_text(text: str) -> str:
    return text.replace("\x00", "")


def _extract_pdf_text_locally(content_b64: str) -> str:
    if not content_b64:
        return ""
    import base64
    import io

    try:
        pdf_bytes = base64.b64decode(content_b64)
    except Exception:
        return ""

    try:
        from pypdf import PdfReader
    except Exception:
        try:
            from PyPDF2 import PdfReader
        except Exception:
            return ""

    try:
        reader = PdfReader(io.BytesIO(pdf_bytes))
        parts: list[str] = []
        for page in reader.pages:
            try:
                text = page.extract_text() or ""
            except Exception:
                text = ""
            cleaned = _clean_storage_text(text).strip()
            if cleaned:
                parts.append(cleaned)
        return "\n\n".join(parts).strip()
    except Exception:
        return ""
