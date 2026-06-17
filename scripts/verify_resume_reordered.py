from __future__ import annotations

from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "sun_qizheng_ai_resume_single_column.docx"


def extract_text() -> str:
    document = Document(DOCX)
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.extend(paragraph.text for paragraph in cell.paragraphs)
    return "\n".join(parts)


if __name__ == "__main__":
    text = extract_text()
    for item in ["主修课程", "自我评价", "专业技能", "项目经历", "基于 Advanced RAG"]:
        print(f"{item}: {item in text}")
    print(f"专业技能位置: {text.find('专业技能')}")
    print(f"项目经历位置: {text.find('项目经历')}")
