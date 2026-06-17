from __future__ import annotations

from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "sun_qizheng_ai_resume_single_column.docx"


def extract_text() -> str:
    document = Document(DOCX)
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.extend(paragraph.text for paragraph in cell.paragraphs)
    return "\n".join(parts)


if __name__ == "__main__":
    text = extract_text()
    checks = [
        "Advanced RAG",
        "RAG Pipeline",
        "Query Rewrite",
        "Multi-Query Expansion",
        "Rerank",
        "Parent-Child",
        "Neighbor Window",
        "recall@k",
        "precision@k",
        "citation_hit",
    ]
    for item in checks:
        print(f"{item}: {item in text}")
