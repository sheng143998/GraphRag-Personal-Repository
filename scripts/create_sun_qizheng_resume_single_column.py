from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "sun_qizheng_ai_resume_single_column.docx"

ACCENT = RGBColor(36, 84, 132)
DARK = RGBColor(31, 45, 61)
BODY = RGBColor(45, 45, 45)
MUTED = RGBColor(92, 105, 117)


def add_run(paragraph, text: str, *, bold=False, size: float | None = None, color=None):
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return run


def add_para(doc, *, before=0, after=2, align=None, line=1.0):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line
    if align:
        paragraph.alignment = align
    return paragraph


def add_heading(doc, text: str) -> None:
    paragraph = add_para(doc, before=7, after=3)
    add_run(paragraph, text, bold=True, size=11.5, color=ACCENT)
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "245484")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_bullet(doc, text: str) -> None:
    paragraph = add_para(doc, after=1, line=1.0)
    paragraph.paragraph_format.left_indent = Cm(0.45)
    paragraph.paragraph_format.first_line_indent = Cm(-0.25)
    add_run(paragraph, "• ", color=ACCENT)
    add_run(paragraph, text, color=BODY)


def add_labeled(doc, label: str, body: str) -> None:
    paragraph = add_para(doc, after=1, line=1.0)
    add_run(paragraph, f"{label}：", bold=True, color=DARK)
    add_run(paragraph, body, color=BODY)


def add_project(doc, name: str, role_date: str, stack: str, intro: str, bullets: list[str]) -> None:
    paragraph = add_para(doc, before=3, after=1)
    add_run(paragraph, name, bold=True, size=10.2, color=DARK)
    add_run(paragraph, f"    {role_date}", size=9, color=MUTED)
    add_labeled(doc, "技术栈", stack)
    add_labeled(doc, "项目简介", intro)
    for bullet in bullets:
        add_bullet(doc, bullet)


def build() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.35)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(1.55)
    section.right_margin = Cm(1.55)

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(9)
    normal.font.color.rgb = BODY

    p = add_para(doc, after=1, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_run(p, "孙齐政", bold=True, size=22, color=ACCENT)
    p = add_para(doc, after=2, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_run(p, "电话：13283766327  |  邮箱：1439986513@qq.com", size=9.2, color=DARK)
    p = add_para(doc, after=5, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_run(p, "求职意向：AI 应用开发工程师  |  期望薪资：面议", bold=True, size=9.3, color=DARK)

    add_heading(doc, "教育背景")
    p = add_para(doc, after=1)
    add_run(p, "中南林业科技大学", bold=True, color=DARK)
    add_run(p, "  |  计算机科学与技术（本科）  |  计算机科学与技术专业", color=BODY)

    add_heading(doc, "奖项荣誉")
    add_bullet(doc, "2025 年湖南省大学生程序设计竞赛 —— 铜奖")
    add_bullet(doc, "2025 年 ICPC 国际大学生程序设计竞赛（香港区域赛）—— 优胜奖")

    add_heading(doc, "项目经历")
    add_project(
        doc,
        "基于 Advanced RAG 的本地知识库智能问答系统",
        "核心开发者（RAG + 全栈） | 2025 年至今",
        "LangChain、LangGraph、PostgreSQL(pgvector)、FastAPI、Spring Boot、Vue 3、DashScope OpenAI-compatible API",
        "从 0 到 1 设计并实现面向个人知识库的 Advanced RAG 系统，覆盖文档解析、Chunk 切分、Embedding 入库、混合检索、Query Rewrite、Rerank、上下文压缩与 LLM 生成，并用于评估不同 RAG 策略在真实知识问答场景下的召回质量、引用准确性与答案忠实性。",
        [
            "构建端到端 RAG Pipeline：实现文档上传、MinerU / Docx / Markdown 解析、文本清洗、Chunk 切分、Embedding 生成、pgvector 入库、召回、上下文组装与 LLM 生成的完整链路。",
            "实现 Advanced RAG 策略引擎：支持 basic-rag、hybrid-rerank、metadata-filter、parent-child、advanced-rag 等策略，通过策略模式统一调度，便于对比不同检索增强方案。",
            "强化 Query Understanding：接入 LLM 进行 Query Rewrite 与 Multi-Query Expansion，将用户原始问题改写为更适合检索的查询，并扩展同义表达与相关子问题。",
            "构建混合检索与重排序流程：基于 PostgreSQL pgvector 实现语义向量召回，结合 PostgreSQL 全文检索实现关键词召回，并接入 Rerank 模型对候选 Chunk 二次排序。",
            "设计上下文增强机制：针对 Chunk 过短导致的语义不完整问题，引入 Parent-Child / Neighbor Window 补全机制，并实现 Query-Aware 句子级压缩，减少无关上下文对 LLM 的干扰。",
            "建立评估闭环：实现评测集管理与实验评估流程，支持 recall@k、precision@k、MRR、citation_hit 等指标，用于分析不同策略的召回与引用效果。",
        ],
    )
    add_project(
        doc,
        "智能化本地生活服务平台",
        "核心开发者 | 2025.04 - 2025.06",
        "Spring Boot、Spring AI Alibaba、MySQL、Redis Stack、RabbitMQ、Resilience4j",
        "面向本地生活场景的 AI 赋能全链路交易平台，整合 RAG 智能客服与高并发秒杀架构。",
        [
            "AI 智能体与 RAG：基于 Spring AI + MCP 协议构建智能客服，集成 Tool Calling 实现“查单-售后”自动闭环；利用 Redis Stack 搭建 RAG 知识库。",
            "高并发架构设计：设计 Redis + Caffeine 二级缓存架构，利用 Guava 令牌桶限流；通过 Redisson 分布式锁 + Lua 脚本保障库存不超卖。",
            "链路优化与安全：引入 RabbitMQ 对下单主链路进行异步解耦，接口平均响应时间下降至 120ms；基于 JWT + Redis 实现无状态鉴权与黑名单管理。",
        ],
    )

    add_heading(doc, "专业技能")
    add_labeled(doc, "编程语言", "熟练掌握 Java、Python；具备扎实的算法基础与代码调试能力。")
    add_labeled(doc, "AI/大模型技术", "深入理解 RAG 架构，熟悉 LangChain、LangGraph；具备 Prompt Engineering 及 LLM/Embedding/Rerank 模型接入调优经验。")
    add_labeled(doc, "后端框架", "熟悉 Spring Boot、FastAPI，能够独立完成全栈项目的开发与部署。")
    add_labeled(doc, "数据库", "熟悉 MySQL、PostgreSQL（含 pgvector 向量扩展）及 Redis，具备混合检索（向量+全文）实战经验。")
    add_labeled(doc, "工具与环境", "熟悉 Docker 容器化技术，熟练使用 Git 进行版本控制。")

    doc.core_properties.title = "孙齐政_AI应用开发工程师_简历_无侧边栏"
    doc.core_properties.author = "孙齐政"
    doc.save(OUT)


if __name__ == "__main__":
    build()
    print(OUT)
