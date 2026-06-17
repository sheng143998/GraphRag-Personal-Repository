from __future__ import annotations

from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "sun_qizheng_ai_resume.docx"


def extract_text() -> str:
    document = Document(DOCX)
    parts: list[str] = []
    for paragraph in document.paragraphs:
        parts.append(paragraph.text)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    parts.append(paragraph.text)
    return "\n".join(parts)


if __name__ == "__main__":
    text = extract_text()
    checks = [
        "孙齐政",
        "13283766327",
        "AI 应用开发工程师",
        "基于 Advanced RAG",
        "智能化本地生活服务平台",
        "自我评价",
    ]
    for item in checks:
        print(f"{item}: {item in text}")
