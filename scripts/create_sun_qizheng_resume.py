from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
ASCII_OUT = ROOT / "sun_qizheng_ai_resume.docx"

ACCENT = RGBColor(36, 84, 132)
DARK = RGBColor(31, 45, 61)
MUTED = RGBColor(90, 100, 110)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_no_borders(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "nil")
        borders.append(tag)
    tc_pr.append(borders)


def add_run(paragraph, text: str, *, bold=False, size: float | None = None, color=None):
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return run


def add_para(container, *, before=0, after=1, align=None, line=1.0):
    paragraph = container.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line
    if align:
        paragraph.alignment = align
    return paragraph


def add_heading(container, text: str) -> None:
    paragraph = add_para(container, before=6, after=2)
    add_run(paragraph, text, bold=True, size=10.5, color=ACCENT)
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "245484")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_bullet(container, text: str) -> None:
    paragraph = add_para(container, after=1, line=1.0)
    paragraph.paragraph_format.left_indent = Cm(0.28)
    paragraph.paragraph_format.first_line_indent = Cm(-0.18)
    add_run(paragraph, "• ", color=ACCENT)
    add_run(paragraph, text)


def add_labeled(container, label: str, body: str) -> None:
    paragraph = add_para(container, after=1, line=1.0)
    add_run(paragraph, f"{label}：", bold=True, color=DARK)
    add_run(paragraph, body)


def add_project(container, name: str, role_date: str, stack: str, intro: str, items: list[str]) -> None:
    paragraph = add_para(container, before=3, after=0)
    add_run(paragraph, name, bold=True, size=9.7, color=DARK)
    add_run(paragraph, f"    {role_date}", size=8.6, color=MUTED)
    add_labeled(container, "技术栈", stack)
    add_labeled(container, "项目简介", intro)
    for item in items:
        add_bullet(container, item)


def build_resume() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.25)
    section.bottom_margin = Cm(1.15)
    section.left_margin = Cm(1.25)
    section.right_margin = Cm(1.25)

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(8.8)
    normal.font.color.rgb = RGBColor(45, 45, 45)

    layout = doc.add_table(rows=1, cols=2)
    layout.alignment = WD_TABLE_ALIGNMENT.CENTER
    layout.autofit = False
    layout.columns[0].width = Cm(5.0)
    layout.columns[1].width = Cm(12.5)
    left, right = layout.rows[0].cells
    for cell in (left, right):
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        set_cell_margins(cell, top=120, start=150, bottom=120, end=150)
        set_no_borders(cell)
    shade_cell(left, "EAF1F8")

    paragraph = add_para(left, after=0, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_run(paragraph, "孙齐政", bold=True, size=21, color=ACCENT)
    paragraph = add_para(left, after=7, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_run(paragraph, "AI 应用开发工程师", bold=True, size=10, color=DARK)

    add_heading(left, "联系方式")
    add_labeled(left, "电话", "13283766327")
    add_labeled(left, "邮箱", "1439986513@qq.com")
    add_labeled(left, "求职意向", "AI 应用开发工程师")
    add_labeled(left, "期望薪资", "面议")

    add_heading(left, "专业技能")
    add_bullet(left, "Java、Python，具备扎实算法基础与调试能力")
    add_bullet(left, "RAG、LangChain、LangGraph、Prompt Engineering")
    add_bullet(left, "Spring Boot、FastAPI、Vue 3 / TypeScript")
    add_bullet(left, "MySQL、PostgreSQL / pgvector、Redis")
    add_bullet(left, "Docker、Git，熟悉项目部署与版本管理")

    add_heading(left, "奖项荣誉")
    add_bullet(left, "2025 年湖南省大学生程序设计竞赛铜奖")
    add_bullet(left, "2025 年 ICPC 国际大学生程序设计竞赛香港区域赛优胜奖")

    add_heading(right, "教育背景")
    paragraph = add_para(right, after=1)
    add_run(paragraph, "中南林业科技大学", bold=True, color=DARK)
    add_run(paragraph, " | 计算机科学与技术（本科） | 计算机科学与技术专业")
    add_labeled(
        right,
        "主修课程",
        "数据结构、算法设计与分析、操作系统、计算机网络、数据库原理、软件工程、面向对象程序设计、人工智能导论等",
    )

    add_heading(right, "项目经历")
    add_project(
        right,
        "基于 Advanced RAG 的本地知识库智能问答系统",
        "核心开发者（全栈） | 2025 年至今",
        "Java/Spring Boot、Python/FastAPI、Vue 3/TypeScript、PostgreSQL(pgvector)、LangChain、LangGraph、OpenAI-compatible API",
        "设计并实现三层架构本地知识库 RAG 系统，覆盖前端交互、业务编排与 AI 检索生成能力，支持多策略检索与智能问答。",
        [
            "全栈架构设计：采用 Vue 3 + Spring Boot + FastAPI 三层分离架构，Java 层负责业务编排与持久化，Python 层负责 AI 检索与生成逻辑。",
            "Advanced RAG 多策略引擎：实现 basic-rag、hybrid-rerank、metadata-filter、parent-child、advanced-rag 等策略，集成 Query Rewrite、Multi-Query Expansion 与 Rerank。",
            "混合检索与向量数据库：基于 PostgreSQL pgvector 实现向量相似度检索，结合全文搜索构建混合检索，并支持不同检索倾向配置。",
            "模型适配与稳定性保障：接入阿里百炼 DashScope，实现 OpenAI-compatible adapter，支持自动降级与指数退避重试。",
            "文档解析与上下文增强：实现异步摄入、解析、分块、向量化流程，引入 Neighbor Window 与 Query-Aware 句子级压缩增强答案依据。",
        ],
    )
    add_project(
        right,
        "智能化本地生活服务平台",
        "核心开发者 | 2025.04 - 2025.06",
        "Spring Boot、Spring AI Alibaba、MySQL、Redis Stack、RabbitMQ、Resilience4j",
        "面向本地生活场景的 AI 赋能全链路交易平台，整合 RAG 智能客服与高并发秒杀架构。",
        [
            "AI 智能体与 RAG：基于 Spring AI + MCP 协议构建智能客服，集成 Tool Calling 实现查单-售后自动闭环；利用 Redis Stack 构建 RAG 知识库。",
            "高并发架构设计：设计 Redis + Caffeine 二级缓存，使用 Guava 令牌桶限流，并通过 Redisson 分布式锁 + Lua 脚本保障库存扣减一致性。",
            "链路优化与安全：引入 RabbitMQ 对下单主链路异步解耦，接口平均响应时间降至 120ms；基于 JWT + Redis 实现无状态鉴权与黑名单管理。",
        ],
    )

    add_heading(right, "自我评价")
    add_bullet(right, "计算机科班出身，具备扎实的数据结构与算法基础，拥有 ICPC / 省赛竞赛经历。")
    add_bullet(right, "对 AI 技术保持高热情，具备从 0 到 1 搭建 Advanced RAG 系统的全栈实战经验。")
    add_bullet(right, "熟悉 AI 应用开发完整链路，能够从业务问题出发设计并落地可迭代的技术方案。")

    doc.core_properties.title = "孙齐政_AI应用开发工程师_简历"
    doc.core_properties.author = "孙齐政"
    doc.core_properties.subject = "AI 应用开发工程师简历"
    doc.save(ASCII_OUT)


if __name__ == "__main__":
    build_resume()
    print(ASCII_OUT)
